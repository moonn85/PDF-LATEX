import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from pathlib import Path
import PyPDF2
import os
import fitz  
import re
import zipfile
import shutil
from datetime import datetime
import docx  
import unicodedata
from docx.shared import Pt
from docx.oxml import OxmlElement
from xml.etree import ElementTree

class PDFToLatexConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("PDF to LaTeX Converter")
        self.root.geometry("600x350")
        self.root.resizable(False, False)
        self.root.configure(bg='#f0f0f0')
        
        # Create main frame
        main_frame = tk.Frame(root, bg='#f0f0f0', padx=20, pady=20)
        main_frame.pack(expand=True, fill='both')
        
        # Title
        title_label = tk.Label(
            main_frame, 
            text="PDF to LaTeX Converter",
            font=("Arial", 16, "bold"),
            bg='#f0f0f0'
        )
        title_label.pack(pady=10)
        
        # File selection frame
        file_frame = tk.LabelFrame(
            main_frame,
            text="File Selection",
            font=("Arial", 10),
            bg='#f0f0f0',
            padx=10,
            pady=10
        )
        file_frame.pack(fill='x', pady=10)
        
        self.selected_file = tk.StringVar()
        tk.Entry(
            file_frame,
            textvariable=self.selected_file,
            width=50,
            font=("Arial", 10)
        ).pack(side='left', padx=5)
        
        browse_btn = tk.Button(
            file_frame,
            text="Browse",
            command=self.browse_file,
            font=("Arial", 10),
            bg='#2196F3',
            fg='white',
            padx=10
        )
        browse_btn.pack(side='left')
        
        # Progress frame
        progress_frame = tk.LabelFrame(
            main_frame,
            text="Conversion Progress",
            font=("Arial", 10),
            bg='#f0f0f0',
            padx=10,
            pady=10
        )
        progress_frame.pack(fill='x', pady=10)
        
        self.progress_var = tk.DoubleVar()
        self.progress_label = tk.Label(
            progress_frame,
            text="0%",
            font=("Arial", 10),
            bg='#f0f0f0'
        )
        self.progress_label.pack()
        
        self.progress_bar = ttk.Progressbar(
            progress_frame,
            variable=self.progress_var,
            maximum=100,
            mode='determinate',
            length=500
        )
        self.progress_bar.pack(pady=5)
        
        # Convert button
        self.convert_btn = tk.Button(
            main_frame,
            text="Convert to LaTeX",
            command=self.convert_document,
            font=("Arial", 12, "bold"),
            bg='#4CAF50',
            fg='white',
            padx=20,
            pady=10
        )
        self.convert_btn.pack(pady=20)

    def browse_file(self):
        filename = filedialog.askopenfilename(
            title="Select File",
            filetypes=[
                ("Supported files", "*.pdf;*.doc;*.docx"),
                ("PDF files", "*.pdf"),
                ("Word files", "*.doc;*.docx"),
            ]
        )
        if filename:
            self.selected_file.set(filename)

    def update_progress(self, current, total):
        progress = (current / total) * 100
        self.progress_var.set(progress)
        self.progress_label.config(text=f"{progress:.1f}%")
        self.root.update()

    def extract_images(self, pdf_path, output_folder):
        """Trích xuất hình ảnh từ PDF"""
        pdf_document = fitz.open(pdf_path)
        images_info = []
        
        for page_number in range(pdf_document.page_count):
            page = pdf_document[page_number]
            image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                base_image = pdf_document.extract_image(xref)
                image_ext = base_image["ext"]
                image_name = f"image_p{page_number + 1}_{img_index + 1}.{image_ext}"
                image_path = os.path.join(output_folder, image_name)
                
                with open(image_path, "wb") as image_file:
                    image_file.write(base_image["image"])
                
                images_info.append({
                    'path': image_name,
                    'page': page_number
                })
        
        return images_info

    def extract_docx_images(self, doc, output_folder):
        """Trích xuất hình ảnh từ DOCX"""
        images_info = []
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.reltype:
                image_data = rel.target_part.blob
                image_ext = rel.target_ref.split('.')[-1]
                image_name = f"image_{i + 1}.{image_ext}"
                image_path = os.path.join(output_folder, image_name)
                
                with open(image_path, "wb") as img_file:
                    img_file.write(image_data)
                
                images_info.append({
                    'path': image_name,
                    'rel_id': rel.rId
                })
        return images_info

    def create_latex_document(self):
        """Tạo template LaTeX với pdfLaTeX để hỗ trợ tiếng Việt"""
        return r"""% !TEX encoding = UTF-8

\documentclass[12pt,a4paper]{article}

% Hỗ trợ tiếng Việt
\usepackage[utf8]{inputenc}
\usepackage[T5]{fontenc}
\usepackage{vntex}

% Các gói cơ bản
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{float}
\usepackage{tabularx}
\usepackage{booktabs}
\usepackage{multirow}
\usepackage{hyperref}
\usepackage{listings}
\usepackage{xcolor}
\usepackage{caption}

% Cấu hình font chữ tiếng Việt
\usepackage{times}

% Cấu hình trang
\geometry{
    a4paper,
    left=2.5cm,
    right=2.5cm,
    top=2.5cm,
    bottom=2.5cm
}

% Đường dẫn hình ảnh
\graphicspath{{./images/}}

% Cấu hình hyperref
\hypersetup{
    unicode=true,
    colorlinks=true,
    linkcolor=blue,
    filecolor=magenta,
    urlcolor=cyan,
}

\begin{document}
"""

    def normalize_vietnamese_text(self, text):
        """Chuẩn hóa văn bản tiếng Việt"""
        # Chuẩn hóa Unicode NFKC thay vì NFC
        text = unicodedata.normalize('NFKC', text)
        text = ' '.join(text.strip().split())
        return text

    def convert_document(self):
        file_path = self.selected_file.get()
        if not file_path:
            messagebox.showerror("Error", "Please select a file first!")
            return
            
        # Tạo thư mục project
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        project_name = f"latex_project_{timestamp}"
        project_dir = Path(file_path).parent / project_name
        images_dir = project_dir / "images"
        os.makedirs(images_dir, exist_ok=True)

        try:
            self.convert_btn.config(state='disabled')
            self.progress_var.set(0)
            
            # Tạo single tex file
            tex_path = project_dir / "main.tex"
            
            with open(tex_path, 'w', encoding='utf-8') as tex_file:
                # Viết phần header với UTF-8 BOM để đảm bảo encoding
                tex_file.write('\ufeff')  # UTF-8 BOM
                tex_file.write(self.create_latex_document())
                
                # Chuyển đổi nội dung theo loại file
                file_ext = Path(file_path).suffix.lower()
                if file_ext == '.pdf':
                    images_info = self.extract_images(file_path, str(images_dir))
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        total_pages = len(pdf_reader.pages)
                        
                        for i, page in enumerate(pdf_reader.pages):
                            text = page.extract_text()
                            text = self.normalize_vietnamese_text(text)
                            text = self.process_tables(text)
                            
                            paragraphs = text.split('\n\n')
                            for para in paragraphs:
                                if para.strip():
                                    para = self.escape_latex(para)
                                    tex_file.write(f"{para}\n\n")
                            
                            # Chèn hình ảnh
                            page_images = [img for img in images_info if img['page'] == i]
                            for img in page_images:
                                tex_file.write(f"""
\\begin{{figure}}[H]
    \\centering
    \\includegraphics[width=0.8\\textwidth]{{images/{img['path']}}}
    \\caption{{}}
    \\label{{fig:{img['path']}}}
\\end{{figure}}
""")
                            tex_file.write("\\newpage\n")
                            self.update_progress(i + 1, total_pages)
                
                elif file_ext in ['.doc', '.docx']:
                    try:
                        doc = docx.Document(file_path)
                        images_info = self.extract_docx_images(doc, str(images_dir))
                        total_items = len(doc.paragraphs) + len(doc.tables)
                        current_item = 0

                        # Process document content
                        for paragraph in doc.paragraphs:
                            self.process_docx_paragraph(paragraph, tex_file)
                            self.process_docx_image_runs(paragraph, images_info, tex_file)
                            current_item += 1
                            self.update_progress(current_item, total_items)

                        # Process tables separately
                        for table in doc.tables:
                            self.process_docx_table(table, tex_file)
                            current_item += 1
                            self.update_progress(current_item, total_items)

                    except Exception as e:
                        print(f"Error processing DOCX: {str(e)}")
                        raise
                
                # Kết thúc document
                tex_file.write("\\end{document}")

            # Tạo README
            readme_content = """LaTeX Project Files
=================
1. main.tex: Main LaTeX file (compile this)
2. images/: Contains all extracted images

Instructions:
1. Keep all files in the same directory structure
2. Compile main.tex with your LaTeX compiler
"""
            with open(project_dir / "README.txt", 'w', encoding='utf-8') as readme_file:
                readme_file.write(readme_content)

            # Nén project
            zip_path = str(project_dir.parent / f"{project_name}.zip")
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                for root, _, files in os.walk(project_dir):
                    for file in files:
                        file_path = Path(root) / file
                        arcname = file_path.relative_to(project_dir)
                        zipf.write(file_path, arcname)

            # Dọn dẹp
            shutil.rmtree(project_dir)
            
            messagebox.showinfo(
                "Success", 
                f"Conversion completed successfully!\n"
                f"Upload this file to Overleaf:\n{zip_path}"
            )
            
        except Exception as e:
            messagebox.showerror("Error", f"Conversion failed: {str(e)}")
        finally:
            self.convert_btn.config(state='normal')
            self.progress_var.set(100)
            self.progress_label.config(text="100%")

    def escape_latex(self, text):
        """Escape các ký tự đặc biệt của LaTeX và xử lý Unicode"""
        text = self.normalize_vietnamese_text(text)
        chars = {
            '\\': '\\textbackslash',
            '&': '\\&',
            '%': '\\%',
            '$': '\\$',
            '#': '\\#',
            '_': '\\_',
            '{': '\\{',
            '}': '\\}',
            '~': '\\textasciitilde',
            '^': '\\textasciicircum'
        }
        for char, replacement in chars.items():
            text = text.replace(char, replacement)
        return text

    def process_tables(self, text):
        """Xử lý và chuyển đổi bảng sang định dạng LaTeX"""
        # Tìm và chuyển đổi các bảng đơn giản
        table_pattern = r'([^\n]+\|[^\n]+)'
        tables = re.finditer(table_pattern, text)
        
        for table in tables:
            table_text = table.group(0)
            rows = table_text.split('\n')
            if len(rows) > 1:  # Đảm bảo đây thực sự là bảng
                latex_table = "\\begin{table}[H]\n\\centering\n\\begin{tabularx}{\\textwidth}{"
                
                # Xác định số cột
                cols = len(rows[0].split('|'))
                latex_table += "X" * cols + "}\n\\toprule\n"
                
                # Thêm nội dung bảng
                for i, row in enumerate(rows):
                    cells = [cell.strip() for cell in row.split('|')]
                    latex_table += " & ".join(self.escape_latex(cell) for cell in cells)
                    latex_table += " \\\\\n"
                    if i == 0:  # Thêm đường kẻ sau hàng tiêu đề
                        latex_table += "\\midrule\n"
                
                latex_table += "\\bottomrule\n\\end{tabularx}\n\\end{table}\n\n"
                text = text.replace(table_text, latex_table)
        
        return text

    def process_docx_paragraph(self, paragraph, tex_file):
        """Xử lý đoạn văn từ DOCX"""
        try:
            if not paragraph.text.strip():
                tex_file.write("\n")
                return

            # Xử lý font size
            font_size = None
            for run in paragraph.runs:
                if hasattr(run, 'font') and run.font.size:
                    font_size = run.font.size.pt
                    break

            text = self.escape_latex(paragraph.text)
            
            # Thêm điều chỉnh font size nếu có
            if font_size:
                text = f"{{\\fontsize{{{font_size}}}{{1.2\\baselineskip}}\\selectfont {text}}}"

            # Xử lý căn lề
            if hasattr(paragraph, 'alignment'):
                if paragraph.alignment == 1:  # Center
                    tex_file.write(f"\\begin{{center}}\n{text}\n\\end{{center}}\n")
                elif paragraph.alignment == 2:  # Right
                    tex_file.write(f"\\begin{{flushright}}\n{text}\n\\end{{flushright}}\n")
                else:  # Left or Justified
                    tex_file.write(f"{text}\\par\n")
            else:
                tex_file.write(f"{text}\\par\n")

        except Exception as e:
            print(f"Error in process_docx_paragraph: {str(e)}")
            # Fallback to simple text output
            tex_file.write(f"{paragraph.text}\\par\n")

    def process_docx_image_runs(self, paragraph, images_info, tex_file):
        """Xử lý hình ảnh trong paragraph một cách an toàn"""
        try:
            for run in paragraph.runs:
                # Kiểm tra drawing elements
                drawings = run._element.findall('.//w:drawing', 
                    {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                if drawings:
                    for drawing in drawings:
                        # Tìm blip element chứa hình ảnh
                        blips = drawing.findall('.//a:blip',
                            {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'})
                        for blip in blips:
                            # Lấy relationship ID từ attribute r:embed
                            rel_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if rel_id:
                                img = next((img for img in images_info if img['rel_id'] == rel_id), None)
                                if img:
                                    tex_file.write(f"""
\\begin{{figure}}[H]
    \\centering
    \\includegraphics[width=0.8\\textwidth]{{images/{img['path']}}}
    \\caption{{}}
\\end{{figure}}
""")
        except Exception as e:
            print(f"Warning: Could not process images in paragraph: {str(e)}")

    def process_docx_table(self, table, tex_file):
        """Xử lý bảng từ DOCX"""
        num_cols = len(table.columns)
        tex_file.write("\n\\begin{table}[H]\n\\centering\n")
        tex_file.write("\\begin{tabularx}{\\textwidth}{|" + "X|" * num_cols + "}\n")
        tex_file.write("\\hline\n")

        for row in table.rows:
            cell_texts = []
            for cell in row.cells:
                # Kết hợp tất cả text trong cell, giữ lại xuống dòng
                cell_text = []
                for paragraph in cell.paragraphs:
                    if paragraph.text.strip():
                        cell_text.append(self.escape_latex(paragraph.text.strip()))
                cell_content = " \\\\ ".join(cell_text) if cell_text else " "
                cell_texts.append(cell_content)
            
            tex_file.write(" & ".join(cell_texts) + " \\\\ \\hline\n")

        tex_file.write("\\end{tabularx}\n")
        tex_file.write("\\caption{}\n")  # Caption trống
        tex_file.write("\\end{table}\n\n")

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFToLatexConverter(root)
    root.mainloop()
